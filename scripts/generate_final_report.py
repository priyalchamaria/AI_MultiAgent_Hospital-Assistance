from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from playwright.sync_api import sync_playwright

ROOT_DIR = Path(__file__).resolve().parent.parent
if str(ROOT_DIR) not in sys.path:
    sys.path.insert(0, str(ROOT_DIR))

APP_URL = "http://localhost:8501"
EDGE_PATH = r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe"
ASSET_DIR = ROOT_DIR / "docs" / "report_assets"
REPORT_PATH = ROOT_DIR / "docs" / "CareBridge_AI_Final_Project_Report.docx"


def capture_screenshots() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    screenshots = {
        "chat_home": ASSET_DIR / "01_chat_home.png",
        "visitor_answer": ASSET_DIR / "02_visitor_answer.png",
        "database": ASSET_DIR / "03_database_explorer.png",
        "policies": ASSET_DIR / "04_policy_library.png",
        "validation": ASSET_DIR / "05_validation_results.png",
        "about": ASSET_DIR / "06_about_architecture.png",
    }

    with sync_playwright() as playwright:
        browser = playwright.chromium.launch(
            headless=True,
            executable_path=EDGE_PATH,
            args=["--no-sandbox"],
        )
        page = browser.new_page(viewport={"width": 1440, "height": 950}, device_scale_factor=1)
        page.goto(APP_URL, wait_until="networkidle", timeout=60000)
        page.screenshot(path=screenshots["chat_home"], full_page=True)

        page.get_by_label("Or type your own question").fill("What are the visiting hours?")
        page.get_by_role("button", name="Ask").click()
        page.wait_for_timeout(2500)
        page.screenshot(path=screenshots["visitor_answer"], full_page=True)

        page.get_by_role("tab", name="Database").click()
        page.wait_for_timeout(1200)
        page.screenshot(path=screenshots["database"], full_page=True)

        page.get_by_role("tab", name="Policies").click()
        page.wait_for_timeout(1200)
        page.screenshot(path=screenshots["policies"], full_page=True)

        page.get_by_role("tab", name="Validation").click()
        page.wait_for_timeout(1200)
        page.screenshot(path=screenshots["validation"], full_page=True)

        page.get_by_role("tab", name="About").click()
        page.wait_for_timeout(1200)
        page.screenshot(path=screenshots["about"], full_page=True)
        browser.close()

    return screenshots


def add_heading(document: Document, text: str, level: int = 1) -> None:
    heading = document.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Number")


def add_code_block(document: Document, lines: list[str]) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_screenshot(document: Document, path: Path, caption: str) -> None:
    if path.exists():
        document.add_picture(str(path), width=Inches(6.4))
        caption_paragraph = document.add_paragraph(caption)
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_paragraph.runs[0].italic = True


def load_validation_summary() -> dict:
    path = ROOT_DIR / "evaluation" / "results.json"
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def build_report(screenshots: dict[str, Path]) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    title = document.add_heading("CareBridge AI: Multi-Agent Hospital Assistant", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph("Final Project Submission Report")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(
        "This report explains the project objective, architecture, implementation, validation, screenshots, "
        "GitHub submission process, and evaluator run instructions."
    )

    add_heading(document, "1. Project Overview")
    document.add_paragraph(
        "CareBridge AI is an AI-powered multi-agent application that allows hospital staff to query "
        "synthetic patient records and synthetic hospital policy documents using plain English. "
        "The system uses an Orchestrator Agent to route each question to the SQL Agent, RAG Agent, "
        "both agents, or an unsupported-query response."
    )

    add_heading(document, "2. Problem Statement")
    document.add_paragraph(
        "Hospital staff often need information from both structured databases and unstructured policy documents. "
        "Traditional systems require separate searches. CareBridge AI combines both information sources into one "
        "conversational interface while keeping the workflow explainable and secure."
    )

    add_heading(document, "3. Objectives")
    add_bullets(
        document,
        [
            "Convert synthetic healthcare CSV data into a clean SQLite patient database.",
            "Generate synthetic hospital policy documents for RAG-based retrieval.",
            "Build a SQL Agent for structured patient-record questions.",
            "Build a RAG Agent for policy-document questions with citations.",
            "Use an Orchestrator Agent to classify and route queries.",
            "Add SQL safety validation, query logging, feedback, role-based access display, and evaluation metrics.",
            "Provide a professional Streamlit interface for easy demonstration.",
        ],
    )

    add_heading(document, "4. System Architecture")
    document.add_paragraph(
        "Core flow: User Query -> Orchestrator Agent -> SQL Agent / RAG Agent / Both Agents / Unsupported -> "
        "Safety Layer -> Unified Response -> Query Logs and Feedback."
    )
    add_bullets(
        document,
        [
            "Orchestrator Agent: classifies each query and records routing reason.",
            "SQL Agent: generates, validates, executes, and formats SQLite queries.",
            "RAG Agent: loads policy documents, chunks them by policy section, retrieves relevant chunks, and returns grounded answers.",
            "Services Layer: query classification, SQL validation, logging, and response formatting.",
            "Frontend: Streamlit dashboard with chat, database explorer, policy library, activity logs, validation, and architecture overview.",
        ],
    )
    add_screenshot(document, screenshots["about"], "Figure 1: About page showing the architecture flow.")

    add_heading(document, "5. Dataset and Database")
    document.add_paragraph(
        "The patient dataset is synthetic. During preparation, the CSV is cleaned and converted into a SQLite database. "
        "The pipeline normalizes column names, converts dates, removes duplicates, checks missing values, generates patient IDs, "
        "creates the patients table, and adds indexes for common query columns."
    )
    add_bullets(
        document,
        [
            "Database file: data/hospital.db",
            "Main table: patients",
            "Indexed columns: medical_condition, doctor, hospital, admission_date, admission_type, test_results",
            "Example fields: patient_id, name, age, gender, condition, doctor, hospital, billing amount, medication, test results",
        ],
    )
    add_screenshot(document, screenshots["database"], "Figure 2: Database Explorer with patient-data summaries.")

    add_heading(document, "6. Policy Documents and RAG")
    document.add_paragraph(
        "The policy corpus contains synthetic hospital policies such as visitor policy, discharge policy, privacy policy, "
        "medication administration policy, infection control policy, and billing policy. Each document contains a title, policy ID, "
        "purpose, scope, responsibilities, procedure, exceptions, escalation process, and review date."
    )
    add_screenshot(document, screenshots["policies"], "Figure 3: Policy Library used as the RAG source corpus.")

    add_heading(document, "7. User Interface")
    document.add_paragraph(
        "The frontend is built in Streamlit because it is easy for evaluators to run locally and provides an interactive dashboard "
        "without a separate frontend/backend setup. The chat page supports a dropdown of demo questions and a custom question input."
    )
    add_screenshot(document, screenshots["chat_home"], "Figure 4: Main chat interface.")
    add_screenshot(document, screenshots["visitor_answer"], "Figure 5: Policy answer with source citation.")

    add_heading(document, "8. Security and Access Control")
    add_bullets(
        document,
        [
            "SQL safety layer allows only SELECT queries.",
            "Unsafe operations such as DROP, DELETE, UPDATE, INSERT, ALTER, TRUNCATE, PRAGMA, and ATTACH are blocked.",
            "Maximum returned rows are controlled.",
            "Role-based access demo changes which patient-table columns are visible for Doctor, Nurse, Billing Staff, and Administrator roles.",
            "Unsupported requests such as weather questions or database-destructive prompts are rejected.",
        ],
    )

    add_heading(document, "9. Evaluation and Testing")
    validation = load_validation_summary()
    if validation:
        document.add_paragraph(
            f"Latest validation summary: routing accuracy = {validation.get('routing_accuracy')}, "
            f"end-to-end success rate = {validation.get('end_to_end_success_rate')}, "
            f"average latency = {validation.get('average_latency')} seconds."
        )
    add_bullets(
        document,
        [
            "Automated tests cover Orchestrator routing, SQL generation, RAG retrieval, and SQL security.",
            "Evaluation metrics include routing accuracy, success rate, unsafe-query blocking, latency, and agent usage distribution.",
            "Feedback buttons collect helpful/not helpful/incorrect signals for future improvement.",
        ],
    )
    add_screenshot(document, screenshots["validation"], "Figure 6: Validation dashboard.")

    add_heading(document, "10. How to Upload Project to GitHub")
    document.add_paragraph("After creating an empty GitHub repository, run these commands from the project folder:")
    add_code_block(
        document,
        [
            "cd C:\\Users\\Priyal\\Desktop\\DS_intern",
            "git init",
            "git add .",
            'git commit -m "Final CareBridge AI project"',
            "git branch -M main",
            "git remote add origin https://github.com/<your-username>/<your-repo-name>.git",
            "git push -u origin main",
        ],
    )
    document.add_paragraph(
        "Before pushing, confirm that .venv, __pycache__, and temporary log files are not included. The .gitignore file already excludes common generated files."
    )

    add_heading(document, "11. How Evaluator Can Download and Run")
    add_numbered(
        document,
        [
            "Open the GitHub repository page.",
            "Click Code and copy the repository URL.",
            "Clone the repository locally using git clone.",
            "Install dependencies from requirements.txt.",
            "Prepare the dataset using the provided archive path or an extracted healthcare_dataset.csv.",
            "Start the Streamlit app and open localhost:8501.",
        ],
    )
    add_code_block(
        document,
        [
            "git clone https://github.com/<your-username>/<your-repo-name>.git",
            "cd <your-repo-name>",
            "pip install -r requirements.txt",
            'python scripts\\prepare_data.py --archive "C:\\path\\to\\archive (1).zip"',
            "python -m streamlit run app.py",
        ],
    )
    document.add_paragraph("Open the application in a browser:")
    add_code_block(document, ["http://localhost:8501"])

    add_heading(document, "12. Demo Questions")
    add_bullets(
        document,
        [
            "How many female patients have diabetes?",
            "What is the average billing amount by insurance provider?",
            "Show the top 5 hospitals with the most cancer patients.",
            "What are the visiting hours?",
            "What is the discharge policy?",
            "How many diabetic patients are there and what is the discharge policy?",
            "What is the weather today?",
        ],
    )

    add_heading(document, "13. Conclusion")
    document.add_paragraph(
        "CareBridge AI demonstrates a practical multi-agent hospital assistant that combines structured patient-record analysis "
        "with policy-document retrieval in one conversational interface. The project includes safety validation, source citations, "
        "role-based display controls, query logging, feedback, tests, and evaluation metrics, making it suitable for final internship submission."
    )

    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(REPORT_PATH)


def main() -> None:
    screenshots = capture_screenshots()
    build_report(screenshots)
    print(f"Generated report: {REPORT_PATH}")
    print(f"Screenshots saved in: {ASSET_DIR}")


if __name__ == "__main__":
    main()
